#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs" / "templates" / "styles" / "word.master-template.docx"


def _require(pkg: str) -> None:
    raise RuntimeError(f"缺少依赖：{pkg}。请先安装 requirements.txt（或 pip install {pkg}）。")


try:
    from docx import Document  # type: ignore
    from docx.shared import Pt  # type: ignore
except Exception:
    _require("python-docx")


def main() -> int:
    TARGET.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # Provide a simple editable starter template
    title = doc.add_heading("Word Master Template", level=0)
    title.runs[0].font.size = Pt(22)
    title.runs[0].bold = True

    doc.add_paragraph("This is the editable DOCX master template for generated outputs.")
    doc.add_paragraph("You can modify styles, page settings, header/footer, and table styles here.")

    doc.save(TARGET)
    print(f"[OK] Created: {TARGET}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

