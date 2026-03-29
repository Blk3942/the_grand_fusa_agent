#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from render_mermaid import render_mmd_to_png
except ImportError:  # pragma: no cover
    render_mmd_to_png = None  # type: ignore[misc, assignment]


ROOT = Path(__file__).resolve().parents[1]

# AUTO_SYNC_BASE_START
SYNC_BASE_WORD_TEMPLATE_MD = "docs/templates/_base/base-template.word.md"
SYNC_BASE_EXCEL_TEMPLATE_MD = "docs/templates/_base/base-template.excel.md"
SYNC_BASE_WORD_HEADINGS = ['Change History（变更历史）', 'General Information（一般信息）', 'Document Purpose（文档目的）', 'Document Scope（文档范围）', 'Input Document（输入文档）', 'References（参考文档）', 'Terms and Abbreviations（术语与缩略语）', 'Terms （术语）', 'Abbreviations（缩略语）']
SYNC_BASE_EXCEL_SHEETS = ['Cover（封面）', 'Document Change History （文档履历）', 'Template Change History （模板履历）', 'General Information（一般信息）', 'Input （输入文档）', 'References（参考文档）']
# AUTO_SYNC_BASE_END
FORMATS_CONFIG = ROOT / "docs" / "templates" / "work-products" / "formats.yaml"
BASE_WORD_TEMPLATE = ROOT / "docs" / "templates" / "_base" / "base.word.yaml"
BASE_EXCEL_TEMPLATE = ROOT / "docs" / "templates" / "_base" / "base.excel.yaml"
EXCEL_STYLES_CONFIG = ROOT / "docs" / "templates" / "styles" / "excel.styles.yaml"
WORD_STYLES_CONFIG = ROOT / "docs" / "templates" / "styles" / "word.styles.yaml"

# Runtime flags (set in main)
include_meta_header = False
include_cover_sheet = True


def _require(pkg: str) -> None:
    raise RuntimeError(
        f"缺少依赖：{pkg}。请先在项目环境中安装 requirements.txt（或手动 pip install {pkg}）。"
    )


try:
    import yaml  # type: ignore
except Exception:  # pragma: no cover
    yaml = None


def load_yaml(path: Path) -> Dict[str, Any]:
    if yaml is None:
        _require("pyyaml")
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def load_formats_config(path: Path) -> Dict[str, Any]:
    if not path.exists():
        raise RuntimeError(f"缺少工作成果输出格式配置文件: {path}")
    return load_yaml(path)


def load_optional_yaml(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {}
    data = load_yaml(path)
    return data if isinstance(data, dict) else {}


def _to_font_kwargs(cfg: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    if "name" in cfg:
        out["name"] = cfg["name"]
    if "size" in cfg:
        out["size"] = cfg["size"]
    if "bold" in cfg:
        out["bold"] = bool(cfg["bold"])
    return out


def _to_alignment_kwargs(cfg: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for k in ["horizontal", "vertical", "wrap_text"]:
        if k in cfg:
            out[k] = cfg[k]
    return out


def normalize_formats(items: Any) -> List[str]:
    if items is None:
        return []
    if isinstance(items, str):
        return [x.strip().lower() for x in items.split(",") if x.strip()]
    if isinstance(items, list):
        return [str(x).strip().lower() for x in items if str(x).strip()]
    raise RuntimeError("formats 配置必须为列表或逗号分隔字符串")


def resolve_formats(work_product_type: str, cli_formats: Optional[str]) -> List[str]:
    cfg = load_formats_config(FORMATS_CONFIG)
    wpf = (cfg or {}).get("work_product_formats", {})
    if work_product_type not in wpf:
        raise RuntimeError(f"未配置 work_product_type 的输出格式: {work_product_type}")

    allowed = set(normalize_formats(wpf[work_product_type].get("allowed_formats")))
    defaults = normalize_formats(wpf[work_product_type].get("default_formats"))
    if not allowed:
        raise RuntimeError(f"{work_product_type} 的 allowed_formats 为空")
    if not defaults:
        raise RuntimeError(f"{work_product_type} 的 default_formats 为空")

    requested = normalize_formats(cli_formats) if cli_formats is not None else defaults
    bad = [f for f in requested if f not in allowed]
    if bad:
        raise RuntimeError(
            f"{work_product_type} 不允许输出格式: {bad}（allowed: {sorted(allowed)}）"
        )
    return requested


def get_path(data: Dict[str, Any], path: str) -> Any:
    cur: Any = data
    for part in path.split("."):
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return None
    return cur


def ensure_list(v: Any) -> List[Any]:
    if v is None:
        return []
    if isinstance(v, list):
        return v
    return [v]


@dataclass(frozen=True)
class DiagramEmbed:
    """单张 UML/Mermaid 图在输出包中的位置与相对引用路径。"""

    placement: str
    title: str
    file_stem: str
    rel_png: str
    rel_mmd: str
    png_path: Path
    mmd_path: Path


def _png_ihdr_pixel_size(path: Path) -> Optional[Tuple[int, int]]:
    """Read PNG width/height from IHDR without Pillow."""
    try:
        with path.open("rb") as f:
            if f.read(8) != b"\x89PNG\r\n\x1a\n":
                return None
            _length = int.from_bytes(f.read(4), "big")
            if f.read(4) != b"IHDR":
                return None
            w = int.from_bytes(f.read(4), "big")
            h = int.from_bytes(f.read(4), "big")
            return (w, h)
    except OSError:
        return None


def _sanitize_diagram_stem(s: str) -> str:
    t = re.sub(r"[^a-zA-Z0-9._-]+", "-", s.strip()).strip("-").lower()
    return t or "diagram"


def prepare_diagram_embeds(
    data: Dict[str, Any],
    uml_dir: Path,
    md_parent: Path,
) -> Dict[str, List[DiagramEmbed]]:
    """
    从 content.diagrams 写入 uml_dir/*.mmd 并渲染 *.png。
    md_parent：生成 .md 所在目录，用于计算 Markdown 内相对路径。
    placement：item_architecture | main_end 等。
    """
    out: Dict[str, List[DiagramEmbed]] = {}
    content = data.get("content", {})
    if not isinstance(content, dict):
        return out
    raw = content.get("diagrams")
    if not raw or not isinstance(raw, list):
        return out
    if render_mmd_to_png is None:
        print("[WARN] 无法导入 render_mmd_to_png，跳过 UML 渲染", file=sys.stderr)
        return out

    uml_dir.mkdir(parents=True, exist_ok=True)
    md_parent = md_parent.resolve()

    for i, item in enumerate(raw):
        if not isinstance(item, dict):
            continue
        pid = str(item.get("id", f"diagram-{i}"))
        title = str(item.get("title", pid))
        placement = str(item.get("placement", "item_architecture")).strip() or "item_architecture"
        stem = item.get("file_stem")
        file_stem = _sanitize_diagram_stem(str(stem)) if stem else _sanitize_diagram_stem(pid)
        mmd_path = uml_dir / f"{file_stem}.mmd"
        png_path = uml_dir / f"{file_stem}.png"

        mermaid_body = item.get("mermaid")
        source_mmd = item.get("source_mmd")
        if mermaid_body is not None and str(mermaid_body).strip():
            mmd_path.write_text(str(mermaid_body).strip() + "\n", encoding="utf-8")
        elif source_mmd and str(source_mmd).strip():
            src = (ROOT / str(source_mmd).strip()).resolve()
            if src.is_file():
                mmd_path.write_text(src.read_text(encoding="utf-8"), encoding="utf-8")
            else:
                print(f"[WARN] diagrams 引用 source_mmd 不存在，已跳过: {source_mmd}", file=sys.stderr)
                continue
        else:
            print(f"[WARN] diagrams 项缺少 mermaid 或 source_mmd，已跳过: id={pid}", file=sys.stderr)
            continue

        if not render_mmd_to_png(mmd_path, png_path):
            print(
                f"[WARN] Mermaid 渲染失败（将仍保留 .mmd，PNG 可能缺失）: {mmd_path.name}",
                file=sys.stderr,
            )

        try:
            rel_png = str(png_path.resolve().relative_to(md_parent)).replace("\\", "/")
            rel_mmd = str(mmd_path.resolve().relative_to(md_parent)).replace("\\", "/")
        except ValueError:
            rel_png = f"uml/{file_stem}.png"
            rel_mmd = f"uml/{file_stem}.mmd"

        de = DiagramEmbed(
            placement=placement,
            title=title,
            file_stem=file_stem,
            rel_png=rel_png,
            rel_mmd=rel_mmd,
            png_path=png_path,
            mmd_path=mmd_path,
        )
        out.setdefault(placement, []).append(de)
    return out


def normalize_base_for_templates(source: Dict[str, Any]) -> Dict[str, Any]:
    """
    兼容层：将旧版 base 字段映射/补齐到新版基础模板期望字段。
    不修改原对象，返回浅拷贝。
    """
    src = dict(source)
    base = dict(src.get("base", {}) or {})

    # Word/Excel 新模板：General Information
    if "document_scope" not in base:
        sb = base.get("scope_boundary", {}) or {}
        sysb = sb.get("system_function_boundary", []) or []
        stages = sb.get("v_model_stage", []) or []
        scope_lines = []
        if sysb:
            scope_lines.append("系统/功能边界：")
            scope_lines.extend([f"- {x}" for x in sysb])
        if stages:
            scope_lines.append("V 模型阶段：")
            scope_lines.extend([f"- {x}" for x in stages])
        base["document_scope"] = "\n".join(scope_lines).strip() if scope_lines else ""

    # Change history：拆分为文档履历/模板履历（若没有则从旧 change_history 派生）
    if "change_history_document" not in base:
        ch = ensure_list(base.get("change_history"))
        base["change_history_document"] = [
            {
                "version": x.get("version", ""),
                "status": x.get("status", base.get("status", "") if isinstance(base.get("status"), str) else ""),
                "author": x.get("author", ""),
                "date": x.get("date", ""),
                "approver": x.get("approver", ""),
                "approve_date": x.get("approve_date", ""),
                "description": x.get("summary", x.get("description", "")),
            }
            for x in ch
            if isinstance(x, dict)
        ]
    if "change_history_template" not in base:
        base["change_history_template"] = []

    # Input documents / references：若未提供，尽量从旧 inputs_references 派生
    if "input_documents" not in base:
        ir = ensure_list(base.get("inputs_references"))
        base["input_documents"] = [
            {"document_no": x.get("id", ""), "file_name": x.get("name", ""), "version": x.get("version", "")}
            for x in ir
            if isinstance(x, dict)
        ]
    if "references" not in base:
        base["references"] = []

    # Terms / abbreviations
    base.setdefault("terms", [])
    base.setdefault("abbreviations", [])

    # Excel template (new) uses one-row table with purpose + scope
    base.setdefault(
        "general_information",
        [{"purpose": str(base.get("purpose", "")), "scope": str(base.get("document_scope", ""))}],
    )

    src["base"] = base
    return src


def md_table(columns: List[str], rows: List[Dict[str, Any]], mapping: Dict[str, str]) -> str:
    lines = []
    lines.append("| " + " | ".join(columns) + " |")
    lines.append("|" + "|".join(["---"] * len(columns)) + "|")
    for row in rows:
        vals = []
        for col in columns:
            key = mapping[col]
            v = row.get(key, "")
            if isinstance(v, list):
                v = ", ".join(str(x) for x in v)
            vals.append(str(v))
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def _md_optional_block(val: Any) -> List[str]:
    """Markdown lines for an optional YAML field (paragraph or bullets)."""
    if val is None or val == "":
        return ["（待填写）"]
    if isinstance(val, list):
        xs = [str(x).strip() for x in val if str(x).strip()]
        return [f"- {x}" for x in xs] if xs else ["（待填写）"]
    return [str(val)]


def _append_use_cases_md(out: List[str], content: Dict[str, Any]) -> None:
    """用例说明：优先 content.use_cases（对齐 template 多用例块）；否则回退为 operating_scenarios 场景表。"""
    ucs = content.get("use_cases")
    if isinstance(ucs, list) and ucs:
        for i, u in enumerate(ucs):
            if not isinstance(u, dict):
                continue
            uid = str(u.get("id", f"UC-{i + 1}"))
            name = str(u.get("name", "")).strip()
            title = f"用例 {uid}" + (f"：{name}" if name else "")
            out.append(f"\n#### {title}\n")
            rows = [
                {"k": "用例名称", "v": name},
                {"k": "主要参与者", "v": u.get("primary_actors", "")},
                {"k": "次要参与者", "v": u.get("secondary_actors", "")},
                {"k": "前置条件", "v": u.get("preconditions", "")},
                {"k": "主成功场景", "v": u.get("main_success", "")},
                {"k": "拓展场景", "v": u.get("extensions", "")},
            ]
            out.append(
                md_table(
                    ["字段", "内容"],
                    rows,
                    {"字段": "k", "内容": "v"},
                )
            )
        return
    out.append(
        md_table(
            ["SCN-ID", "场景描述", "速度/道路/天气等条件", "驾驶员状态", "备注"],
            content.get("operating_scenarios", []) or [],
            {
                "SCN-ID": "id",
                "场景描述": "description",
                "速度/道路/天气等条件": "conditions",
                "驾驶员状态": "driver_state",
                "备注": "note",
            },
        )
    )


def _append_function_partition_md(out: List[str], fp: Any, bi: Dict[str, Any]) -> None:
    if isinstance(fp, list) and fp and isinstance(fp[0], dict):
        out.append(
            md_table(
                ["子功能", "说明", "备注"],
                fp,
                {"子功能": "subfunction", "说明": "description", "备注": "note"},
            )
        )
        return
    if fp:
        out.extend([f"- {x}" for x in ensure_list(fp)])
    else:
        out.append("- （可与「功能简介」「系统边界」一致；以下为当前边界中的在范围要点）")
        ins = bi.get("in_scope", []) or []
        if ins:
            for x in ins[:12]:
                out.append(f"- {x}")
        else:
            out.append("- （待填写）")


def _append_functional_runtime_state_md(out: List[str], frs: Any) -> None:
    if isinstance(frs, dict):
        st = frs.get("states")
        tr = frs.get("transitions")
        if isinstance(st, list) and st:
            out.append("\n")
            out.append(
                md_table(
                    ["状态", "说明", "备注"],
                    st,
                    {"状态": "state", "说明": "description", "备注": "note"},
                )
            )
        if isinstance(tr, list) and tr:
            out.append("\n")
            out.append(
                md_table(
                    ["跳转条件", "说明", "备注"],
                    tr,
                    {"跳转条件": "trigger", "说明": "description", "备注": "note"},
                )
            )
        extra = frs.get("notes") or frs.get("narrative")
        if extra:
            out.append("")
            out.extend(_md_optional_block(extra))
        if not st and not tr and not extra:
            out.append("（待填写）")
        return
    if frs:
        out.extend(_md_optional_block(frs))
    else:
        out.append("（待填写）")


def _append_foreseeable_misuse_md(out: List[str], fm: Any) -> None:
    if isinstance(fm, list) and fm and isinstance(fm[0], dict):
        out.append(
            md_table(
                ["编号", "误用说明", "备注"],
                fm,
                {"编号": "id", "误用说明": "description", "备注": "note"},
            )
        )
        return
    out.extend(_md_optional_block(fm))


def append_item_definition_main_md(
    out: List[str],
    content: Dict[str, Any],
    diagrams_by_placement: Optional[Dict[str, List[DiagramEmbed]]] = None,
) -> None:
    """正文「相关项定义」（含「功能实现」节），对齐 docs/.../item-definition/template.md。"""
    ov = content.get("item_overview", {}) or {}
    bi = content.get("boundary_interfaces", {}) or {}
    da = content.get("driving_automation", {}) or {}
    odd_raw = content.get("odd")
    vss = content.get("vehicle_safety_strategy", {}) or {}
    ad = content.get("alerts_degradation", {}) or {}
    fp = content.get("function_partition")
    frs = content.get("functional_runtime_state")
    fm = content.get("foreseeable_misuse")
    nf = content.get("non_functional", {}) or {}
    dep_o2t = content.get("dependency_other_to_this")
    dep_t2o = content.get("dependency_this_to_other")
    dep_env = content.get("dependency_environment")
    kfm = content.get("known_failure_modes")
    ktc = content.get("known_trigger_conditions")
    osa = content.get("operational_safety_assurance", {}) or {}
    impl = content.get("implementation", {}) or {}

    out.append("\n# 相关项定义\n")

    out.append("\n## 功能定义\n")

    out.append("\n### 功能简介\n")
    out.append(f"- Item 名称：{ov.get('item_name', '')}")
    out.append(f"- 功能目标：{ov.get('functional_objective', '')}")
    out.append("- 非安全目标（如性能/舒适性）：")
    nso = ov.get("non_safety_objectives", []) or []
    if nso:
        for x in nso:
            out.append(f"  - {x}")
    else:
        out.append("  - （待填写）")

    out.append("\n### 驾驶自动化等级\n")
    if isinstance(da, dict) and (da.get("level") or da.get("standard") or da.get("role_phases")):
        if da.get("level"):
            out.append(f"- 驾驶自动化等级：{da.get('level', '')}")
        if da.get("standard"):
            out.append(f"- 参考标准：{da.get('standard', '')}")
        rp = da.get("role_phases") or []
        if rp:
            out.append("\n")
            out.append(
                md_table(
                    ["阶段", "用户的角色", "系统的角色"],
                    rp,
                    {"阶段": "phase", "用户的角色": "user_role", "系统的角色": "system_role"},
                )
            )
    else:
        out.append("（待填写；非驾驶自动化相关项无需填写本节。）")

    out.append("\n### ODD（设计运行范围）\n")
    if isinstance(odd_raw, dict):
        ref = odd_raw.get("reference") or odd_raw.get("description")
        out.extend(_md_optional_block(ref))
    elif odd_raw:
        out.extend(_md_optional_block(odd_raw))
    else:
        out.append(
            "（待填写：引用外部 ODD 文档 / 本节写 N/A 并说明理由 / 不适用则删除并记录在变更说明。）"
        )

    out.append("\n### 整车层面安全策略\n")
    out.append("\n#### 整车安全策略\n")
    out.append("\n##### 法律法规\n")
    out.extend(_md_optional_block(vss.get("laws")))
    out.append("\n##### 驾驶行为\n")
    out.extend(_md_optional_block(vss.get("driving_behavior")))

    out.append("\n#### 整车层面 SOTIF 残余风险接受准则\n")
    out.extend(_md_optional_block(vss.get("sotif_residual")))

    out.append("\n#### 提示报警与降级\n")
    out.append("\n##### 提示报警\n")
    out.extend(_md_optional_block(ad.get("alerts") if ad else vss.get("alerts")))
    out.append("\n##### 降级思路\n")
    out.extend(_md_optional_block(ad.get("rationale") if ad else vss.get("degradation_rationale")))
    out.append("\n##### 降级策略\n")
    rm_states = (ad.get("rm_states") if ad else None) or vss.get("rm_states")
    triggers = (ad.get("triggers") if ad else None) or vss.get("triggers")
    if rm_states:
        out.append("\n**状态说明**\n\n")
        out.append(
            md_table(
                ["状态", "说明", "备注"],
                rm_states,
                {"状态": "state", "说明": "description", "备注": "note"},
            )
        )
    if triggers:
        out.append("\n**跳转条件说明**\n\n")
        out.append(
            md_table(
                ["跳转条件", "说明", "备注"],
                triggers,
                {"跳转条件": "trigger", "说明": "description", "备注": "note"},
            )
        )
    if not rm_states and not triggers:
        out.append("（待填写）")

    out.append("\n## 功能行为与用例\n")
    out.append("\n### 功能划分与简介\n")
    _append_function_partition_md(out, fp, bi)

    out.append("\n### 用例说明\n")
    _append_use_cases_md(out, content)

    out.append("\n## 功能运行状态\n")
    _append_functional_runtime_state_md(out, frs)

    out.append("\n## 合理可预见的误用\n")
    _append_foreseeable_misuse_md(out, fm)

    out.append("\n## 非功能性要求\n")
    out.append("\n### 质量要求\n")
    out.extend(_md_optional_block(nf.get("quality")))
    out.append("\n### 性能与可用性要求\n")
    out.extend(_md_optional_block(nf.get("performance")))

    out.append("\n## 功能依赖 / 假设\n")
    out.append("\n### 其他相关项对本相关项的依赖 / 假设\n")
    out.extend(_md_optional_block(dep_o2t))
    out.append("\n### 本相关项对其他相关项的依赖 / 假设\n")
    out.extend(_md_optional_block(dep_t2o))
    out.append("\n### 环境依赖 / 假设\n")
    out.extend(_md_optional_block(dep_env))
    out.append("\n**依赖与约束（结构化表）**\n")
    out.append(
        md_table(
            ["DEP-ID", "依赖项（系统/传感器/电源/网络）", "约束条件", "失效影响"],
            content.get("dependencies_constraints", []) or [],
            {
                "DEP-ID": "id",
                "依赖项（系统/传感器/电源/网络）": "dependency",
                "约束条件": "constraint",
                "失效影响": "impact_if_failed",
            },
        )
    )

    out.append("\n## 功能层面已知功能局限\n")
    out.append("\n### 已知失效模式\n")
    out.extend(_md_optional_block(kfm))
    out.append("\n### 已知功能触发条件\n")
    out.extend(_md_optional_block(ktc))
    out.append("\n**初始安全关注点（供 HARA 输入）**\n")
    out.append(
        md_table(
            ["CON-ID", "关注点描述", "来源", "严重性初判", "后续流转"],
            content.get("known_failures_initial_concerns", []) or [],
            {
                "CON-ID": "id",
                "关注点描述": "concern",
                "来源": "source",
                "严重性初判": "severity_preliminary",
                "后续流转": "downstream",
            },
        )
    )

    out.append("\n## 运行安全保障\n")
    out.append("\n### 运行阶段安全设计\n")
    if isinstance(osa, dict):
        out.extend(_md_optional_block(osa.get("operational_design")))
    else:
        out.extend(_md_optional_block(osa))

    out.append("\n## 功能实现\n")
    out.append("\n### 功能架构\n")
    out.append("\n#### 架构示意图\n")
    out.extend(_md_optional_block(impl.get("architecture")))
    for d in (diagrams_by_placement or {}).get("item_architecture", []):
        out.append(f"\n**{d.title}**（Mermaid 源：`{d.rel_mmd}`）\n")
        if d.png_path.is_file():
            out.append(f"![{d.title}]({d.rel_png})\n")
        else:
            out.append("*（图未渲染：请配置 Node/mermaid-cli 或可用 Kroki）*\n")
    out.append("\n#### 要素说明\n")
    arch_el = impl.get("architecture_elements")
    if isinstance(arch_el, list) and arch_el:
        out.append(
            md_table(
                ["要素", "说明", "备注"],
                arch_el,
                {"要素": "element", "说明": "description", "备注": "note"},
            )
        )
    else:
        out.append("- 在范围（In Scope）：")
        for x in bi.get("in_scope", []) or []:
            out.append(f"  - {x}")
        if not bi.get("in_scope"):
            out.append("  - （待填写）")
        out.append("- 不在范围（Out of Scope）：")
        for x in bi.get("out_of_scope", []) or []:
            out.append(f"  - {x}")
        if not bi.get("out_of_scope"):
            out.append("  - （待填写）")

    out.append("\n#### 接口说明\n")
    ifb = impl.get("interfaces_brief")
    if isinstance(ifb, list) and ifb:
        out.append(
            md_table(
                ["接口", "说明", "备注"],
                ifb,
                {"接口": "interface", "说明": "description", "备注": "note"},
            )
        )
    else:
        out.append(
            md_table(
                ["IF-ID", "接口对象", "输入/输出", "信号/消息", "正常范围", "异常处理"],
                bi.get("interfaces", []) or [],
                {
                    "IF-ID": "id",
                    "接口对象": "object",
                    "输入/输出": "direction",
                    "信号/消息": "signal",
                    "正常范围": "normal_range",
                    "异常处理": "abnormal_handling",
                },
            )
        )

    out.append("\n### 非功能性要求（实现层）\n")
    out.append("\n需要从概念阶段的非功能性要求进行拆解得出。\n")
    if impl.get("nf_quality"):
        out.append("\n#### 质量要求\n")
        out.extend(_md_optional_block(impl.get("nf_quality")))
    out.append("\n#### 性能与可用性要求\n")
    out.extend(_md_optional_block(impl.get("nf_performance")))

    out.append("\n### 要素层面已知功能局限\n")
    out.append("\n#### 已知失效模式\n")
    out.extend(_md_optional_block(impl.get("element_failure_modes")))
    out.append("\n#### 已知功能触发条件\n")
    out.extend(_md_optional_block(impl.get("element_trigger_conditions")))


def render_md(
    source: Dict[str, Any],
    diagrams_by_placement: Optional[Dict[str, List[DiagramEmbed]]] = None,
) -> str:
    source = normalize_base_for_templates(source)
    meta = source.get("meta", {})
    base = source.get("base", {})
    content = source.get("content", {})
    dmap = diagrams_by_placement or {}

    out: List[str] = []

    # Base sections aligned with latest base-template.word.md
    out.append("## Change History（变更历史）\n")
    out.append("## Document Change History（文档履历）\n")
    out.append(
        md_table(
            ["Version", "Status", "Author", "Date", "Approver", "Approve Date", "Description"],
            base.get("change_history_document", []) or [],
            {
                "Version": "version",
                "Status": "status",
                "Author": "author",
                "Date": "date",
                "Approver": "approver",
                "Approve Date": "approve_date",
                "Description": "description",
            },
        )
    )
    out.append("\n## Template Change History（文档模板履历）\n")
    out.append(
        md_table(
            ["Version", "Status", "Author", "Date", "Approver", "Approve Date", "Description"],
            base.get("change_history_template", []) or [],
            {
                "Version": "version",
                "Status": "status",
                "Author": "author",
                "Date": "date",
                "Approver": "approver",
                "Approve Date": "approve_date",
                "Description": "description",
            },
        )
    )

    out.append("\n# General Information（一般信息）\n")
    out.append("## Document Purpose（文档目的）\n\n" + str(base.get("purpose", "")))
    out.append("\n## Document Scope（文档范围）\n\n" + str(base.get("document_scope", "")))

    out.append("\n## Input Document（输入文档）\n")
    out.append(
        md_table(
            ["Document No.", "File name", "Version"],
            base.get("input_documents", []) or [],
            {"Document No.": "document_no", "File name": "file_name", "Version": "version"},
        )
    )
    out.append("\n## References（参考文档）\n")
    out.append(
        md_table(
            ["Document No.", "File name", "Version"],
            base.get("references", []) or [],
            {"Document No.": "document_no", "File name": "file_name", "Version": "version"},
        )
    )

    out.append("\n# Terms and Abbreviations（术语与缩略语）\n")
    out.append("## Terms （术语）\n")
    out.append(
        md_table(
            ["Term", "Definition", "Source"],
            base.get("terms", []) or [],
            {"Term": "term", "Definition": "definition", "Source": "source"},
        )
    )
    out.append("\n## Abbreviations（缩略语）\n")
    out.append(
        md_table(
            ["Term", "Definition"],
            base.get("abbreviations", []) or [],
            {"Term": "term", "Definition": "definition"},
        )
    )

    # Main content: item-definition + hara（不接「正文内容」总标题，避免与基础模板重复且消除错误章节号 6.x）
    wp_type = meta.get("work_product_type")
    if wp_type == "item-definition":
        append_item_definition_main_md(out, content if isinstance(content, dict) else {}, dmap)
    elif wp_type == "hara":
        scope = (content.get("analysis_scope", {}) or {})
        out.append("\n# 危害分析与风险评估（HARA）\n")
        out.append("## 分析范围\n")
        out.append(f"- 关联 Item：{scope.get('related_item','')}")
        out.append("- 分析边界：")
        for x in scope.get("included_boundary", []) or []:
            out.append(f"  - {x}")
        out.append("- 不覆盖范围：")
        for x in scope.get("excluded_boundary", []) or []:
            out.append(f"  - {x}")

        out.append("\n## 场景与危害事件识别\n")
        out.append(
            md_table(
                ["HE-ID", "运行场景", "malfunction（功能异常）", "危害事件", "潜在后果"],
                content.get("hazards_and_events", []) or [],
                {
                    "HE-ID": "he_id",
                    "运行场景": "scenario",
                    "malfunction（功能异常）": "malfunction",
                    "危害事件": "hazard_event",
                    "潜在后果": "consequence",
                },
            )
        )

        out.append("\n## 风险评估与 ASIL 判定\n")
        out.append(
            md_table(
                ["HE-ID", "S", "E", "C", "ASIL", "判定依据简述"],
                content.get("risk_assessment", []) or [],
                {
                    "HE-ID": "he_id",
                    "S": "s",
                    "E": "e",
                    "C": "c",
                    "ASIL": "asil",
                    "判定依据简述": "rationale",
                },
            )
        )

        out.append("\n## 安全目标（Safety Goals）\n")
        out.append(
            md_table(
                ["SG-ID", "安全目标描述", "来源 HE-ID", "ASIL", "安全状态（Safe State）"],
                content.get("safety_goals", []) or [],
                {
                    "SG-ID": "sg_id",
                    "安全目标描述": "description",
                    "来源 HE-ID": "source_he_id",
                    "ASIL": "asil",
                    "安全状态（Safe State）": "safe_state",
                },
            )
        )

        out.append("\n## 假设、限制与边界条件\n")
        out.append(
            md_table(
                ["类型", "ID", "内容", "影响分析"],
                content.get("assumptions_limits", []) or [],
                {"类型": "type", "ID": "id", "内容": "content", "影响分析": "impact"},
            )
        )

        out.append("\n## 3.6 向下游分配与追溯\n")
        out.append(
            md_table(
                ["SG-ID", "分配目标（FSC/FSR）", "接收方", "验证方法", "证据 ID"],
                content.get("downstream_allocation", []) or [],
                {
                    "SG-ID": "sg_id",
                    "分配目标（FSC/FSR）": "allocation",
                    "接收方": "receiver",
                    "验证方法": "verification_method",
                    "证据 ID": "evidence_id",
                },
            )
        )
        mains = dmap.get("main_end", [])
        if mains:
            out.append("\n## UML 附图\n")
            for d in mains:
                out.append(f"\n### {d.title}\n")
                out.append(f"Mermaid 源：`{d.rel_mmd}`\n")
                if d.png_path.is_file():
                    out.append(f"![{d.title}]({d.rel_png})\n")
                else:
                    out.append("*（图未渲染）*\n")
    else:
        out.append("> 当前脚本仅实现 item-definition 的正文渲染，其他类型可按同样结构扩展。")

    out.append("\n# 可追溯性（Traceability）\n")
    out.append(
        md_table(
            ["条目 ID", "上游输入（父级）", "下游分解（子级）", "验证方法", "证据 ID"],
            base.get("traceability", []) or [],
            {"条目 ID": "item_id", "上游输入（父级）": "upstream", "下游分解（子级）": "downstream", "验证方法": "verification_method", "证据 ID": "evidence_id"},
        )
    )

    out.append("\n# 待验证项（Open Verification Items）\n")
    out.append(
        md_table(
            ["OVI ID", "待验证项", "责任人", "截止日期", "状态"],
            base.get("open_verification_items", []) or [],
            {"OVI ID": "id", "待验证项": "item", "责任人": "owner", "截止日期": "due", "状态": "status"},
        )
    )

    out.append("\n# 评审与批准（Review and Approval）\n")
    out.append(
        md_table(
            ["角色", "姓名", "结论", "日期"],
            base.get("review_approval", []) or [],
            {"角色": "role", "姓名": "name", "结论": "conclusion", "日期": "date"},
        )
    )

    out.append("\n# 变更记录（Change History）\n")
    out.append(
        md_table(
            ["版本", "日期", "变更摘要", "变更原因", "作者"],
            base.get("change_history", []) or [],
            {"版本": "version", "日期": "date", "变更摘要": "summary", "变更原因": "reason", "作者": "author"},
        )
    )

    out.append("\n# 备注与边界声明\n")
    for x in base.get("notes_boundary", []) or []:
        out.append(f"- {x}")

    return "\n".join(out).strip() + "\n"


def render_xlsx(source: Dict[str, Any], out_path: Path) -> None:
    source = normalize_base_for_templates(source)
    try:
        from openpyxl import Workbook  # type: ignore
        from openpyxl.styles import Alignment, Font  # type: ignore
        from openpyxl.styles.borders import Border, Side  # type: ignore
        from openpyxl.utils import column_index_from_string, get_column_letter  # type: ignore
    except Exception:  # pragma: no cover
        _require("openpyxl")

    wb = Workbook()
    # remove default sheet
    default = wb.active
    wb.remove(default)

    meta = source.get("meta", {})
    base = source.get("base", {})

    # Base workbook is driven by template spec
    base_spec = load_yaml(BASE_EXCEL_TEMPLATE)
    excel_styles = load_optional_yaml(EXCEL_STYLES_CONFIG)

    def apply_cell_style(cell: Any, font_cfg: Dict[str, Any], align_cfg: Dict[str, Any]) -> None:
        if font_cfg:
            cell.font = Font(**_to_font_kwargs(font_cfg))
        if align_cfg:
            cell.alignment = Alignment(**_to_alignment_kwargs(align_cfg))

    def build_border(border_cfg: Dict[str, Any]) -> Border | None:
        if not border_cfg:
            return None
        style = border_cfg.get("style", "thin")
        color = border_cfg.get("color", "000000")
        side = Side(style=style, color=color)
        return Border(left=side, right=side, top=side, bottom=side)

    def apply_table_borders(ws: Any, header_border: Border | None, body_border: Border | None, outline_cfg: Dict[str, Any]) -> None:
        if ws.max_row < 1 or ws.max_column < 1:
            return
        # header row
        if header_border:
            for c in ws[1]:
                c.border = header_border
        # body
        if body_border and ws.max_row >= 2:
            for row_idx in range(2, ws.max_row + 1):
                for c in ws[row_idx]:
                    c.border = body_border

        if not outline_cfg or not outline_cfg.get("enabled", False):
            return
        outline_border = build_border(outline_cfg)
        if not outline_border:
            return

        max_r = ws.max_row
        max_c = ws.max_column
        # top and bottom
        for c in ws[1]:
            c.border = outline_border
        for c in ws[max_r]:
            c.border = outline_border
        # left and right
        for r in range(1, max_r + 1):
            ws.cell(row=r, column=1).border = outline_border
            ws.cell(row=r, column=max_c).border = outline_border

    def add_kv_sheet(name: str, rows: List[Dict[str, Any]]) -> None:
        ws = wb.create_sheet(title=name)

        # Special: render Cover as a real cover page
        if name.strip().lower() == "cover" or name.strip().lower().startswith("cover"):
            # Title is expected to be provided by key 'title'
            title_value = ""
            for r in rows:
                if str(r.get("key", "")).strip().lower() == "title":
                    v = get_path(source, r.get("value_path", ""))
                    if isinstance(v, list):
                        v = ", ".join(str(x) for x in v)
                    title_value = "" if v is None else str(v)
                    break

            cover_cfg = excel_styles.get("cover", {}) if isinstance(excel_styles, dict) else {}
            ws.sheet_view.showGridLines = not bool(cover_cfg.get("hide_grid_lines", True))
            merge_range = str(cover_cfg.get("merge_range", "A2:H6"))

            # Apply width/height only to the merge area columns/rows
            m = re.match(r"^([A-Z]+)(\d+):([A-Z]+)(\d+)$", merge_range.strip().upper())
            if m:
                c1, r1, c2, r2 = m.group(1), int(m.group(2)), m.group(3), int(m.group(4))
                col_w = float(cover_cfg.get("column_width", 18))
                row_h = float(cover_cfg.get("row_height", 28))
                for col_idx in range(column_index_from_string(c1), column_index_from_string(c2) + 1):
                    ws.column_dimensions[get_column_letter(col_idx)].width = col_w
                for row_idx in range(min(r1, r2), max(r1, r2) + 1):
                    ws.row_dimensions[row_idx].height = row_h

            ws.merge_cells(merge_range)
            # Top-left cell of merge range
            cell = ws[merge_range.split(":")[0]]
            cell.value = title_value or "（文档标题待填写）"
            apply_cell_style(
                cell,
                (cover_cfg.get("font", {}) if isinstance(cover_cfg, dict) else {}),
                (cover_cfg.get("alignment", {}) if isinstance(cover_cfg, dict) else {}),
            )
            return

        # Default key/value sheet
        ws.append(["key", "value"])
        if not rows:
            # Always keep one empty body row for manual filling
            ws.append(["", ""])
            return
        for r in rows:
            key = r["key"]
            v = get_path(source, r["value_path"])
            if isinstance(v, list):
                v = ", ".join(str(x) for x in v)
            ws.append([key, v])

        table_cfg = excel_styles.get("table", {}) if isinstance(excel_styles, dict) else {}
        header_cfg = table_cfg.get("header", {}) if isinstance(table_cfg, dict) else {}
        body_cfg = table_cfg.get("body", {}) if isinstance(table_cfg, dict) else {}
        outline_cfg = table_cfg.get("outline", {}) if isinstance(table_cfg, dict) else {}
        for c in ws[1]:
            apply_cell_style(c, header_cfg.get("font", {}) if isinstance(header_cfg, dict) else {}, header_cfg.get("alignment", {}) if isinstance(header_cfg, dict) else {})
        for row_idx in range(2, ws.max_row + 1):
            for c in ws[row_idx]:
                apply_cell_style(c, body_cfg.get("font", {}) if isinstance(body_cfg, dict) else {}, body_cfg.get("alignment", {}) if isinstance(body_cfg, dict) else {})
        apply_table_borders(
            ws,
            build_border(header_cfg.get("border", {}) if isinstance(header_cfg, dict) else {}),
            build_border(body_cfg.get("border", {}) if isinstance(body_cfg, dict) else {}),
            outline_cfg if isinstance(outline_cfg, dict) else {},
        )

    def add_table_sheet(name: str, columns: List[str], rows: List[Dict[str, Any]], mapping: Dict[str, str]) -> None:
        ws = wb.create_sheet(title=name)
        ws.append(columns)
        if rows:
            for row in rows:
                out_row = []
                for col in columns:
                    key = mapping[col]
                    v = row.get(key, "")
                    if isinstance(v, list):
                        v = ", ".join(str(x) for x in v)
                    out_row.append(v)
                ws.append(out_row)
        else:
            # Always keep one empty body row for manual filling
            ws.append(["" for _ in columns])

        table_cfg = excel_styles.get("table", {}) if isinstance(excel_styles, dict) else {}
        header_cfg = table_cfg.get("header", {}) if isinstance(table_cfg, dict) else {}
        body_cfg = table_cfg.get("body", {}) if isinstance(table_cfg, dict) else {}
        outline_cfg = table_cfg.get("outline", {}) if isinstance(table_cfg, dict) else {}
        for c in ws[1]:
            apply_cell_style(c, header_cfg.get("font", {}) if isinstance(header_cfg, dict) else {}, header_cfg.get("alignment", {}) if isinstance(header_cfg, dict) else {})
        for row_idx in range(2, ws.max_row + 1):
            for c in ws[row_idx]:
                apply_cell_style(c, body_cfg.get("font", {}) if isinstance(body_cfg, dict) else {}, body_cfg.get("alignment", {}) if isinstance(body_cfg, dict) else {})
        apply_table_borders(
            ws,
            build_border(header_cfg.get("border", {}) if isinstance(header_cfg, dict) else {}),
            build_border(body_cfg.get("border", {}) if isinstance(body_cfg, dict) else {}),
            outline_cfg if isinstance(outline_cfg, dict) else {},
        )

    def text_len(v: Any) -> int:
        if v is None:
            return 0
        s = str(v)
        # for multi-line cells, use longest line
        return max((len(line) for line in s.splitlines()), default=0)

    def auto_fit_sheet_columns(ws: Any, cfg: Dict[str, Any]) -> None:
        if not cfg.get("enabled", False):
            return
        # Skip Cover (fixed layout)
        t = str(ws.title).strip().lower()
        if t == "cover" or t.startswith("cover"):
            return
        min_w = float(cfg.get("min_width", 10))
        max_w = float(cfg.get("max_width", 60))
        padding = float(cfg.get("padding", 2))

        if ws.max_column < 1 or ws.max_row < 1:
            return

        for col_idx in range(1, ws.max_column + 1):
            max_len = 0
            for row_idx in range(1, ws.max_row + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                max_len = max(max_len, text_len(cell.value))
            # very simple approximation: character count + padding
            width = min(max_w, max(min_w, max_len + padding))
            col_letter = ws.cell(row=1, column=col_idx).column_letter
            ws.column_dimensions[col_letter].width = width

    sheets = (((base_spec or {}).get("workbook", {}) or {}).get("sheets", [])) or []
    for sh in sheets:
        name = sh.get("name")
        kind = sh.get("kind")
        if not name or not kind:
            continue
        if kind == "key_value":
            add_kv_sheet(name, sh.get("rows", []) or [])
        elif kind == "table":
            cols = sh.get("columns", []) or []
            src_path = sh.get("source", "")
            rows = get_path(source, src_path) if src_path else []
            if rows is None:
                rows = []
            add_table_sheet(name, cols, rows or [], sh.get("mapping", {}) or {})
        else:
            # unknown kind: skip
            continue

    # Work-product specific sheet(s)
    if meta.get("work_product_type") == "item-definition":
        c = source.get("content", {})
        ov = c.get("item_overview", {})
        bi = c.get("boundary_interfaces", {})

        ws = wb.create_sheet(title="Item")
        ws.append(["Field", "Value"])
        ws.append(["Item 名称", ov.get("item_name", "")])
        ws.append(["功能目标", ov.get("functional_objective", "")])
        ws.append(["非安全目标", ", ".join(ov.get("non_safety_objectives", []) or [])])
        ws.append(["In Scope", "; ".join(bi.get("in_scope", []) or [])])
        ws.append(["Out of Scope", "; ".join(bi.get("out_of_scope", []) or [])])
        da = c.get("driving_automation", {}) or {}
        odd_raw = c.get("odd")
        odd_str = ""
        if isinstance(odd_raw, dict):
            odd_str = str(odd_raw.get("reference") or odd_raw.get("description") or "")
        elif odd_raw:
            odd_str = str(odd_raw)
        ws.append(["ODD（引用/说明）", odd_str])
        if isinstance(da, dict):
            ws.append(["驾驶自动化等级", str(da.get("level", "") or "")])
            ws.append(["自动化参考标准", str(da.get("standard", "") or "")])
        rp = da.get("role_phases", []) if isinstance(da, dict) else []
        if rp:
            add_table_sheet(
                "AutoRoles",
                ["阶段", "用户的角色", "系统的角色"],
                rp,
                {"阶段": "phase", "用户的角色": "user_role", "系统的角色": "system_role"},
            )

        add_table_sheet(
            "Interfaces",
            ["IF-ID", "接口对象", "输入/输出", "信号/消息", "正常范围", "异常处理"],
            bi.get("interfaces", []) or [],
            {
                "IF-ID": "id",
                "接口对象": "object",
                "输入/输出": "direction",
                "信号/消息": "signal",
                "正常范围": "normal_range",
                "异常处理": "abnormal_handling",
            },
        )
        add_table_sheet(
            "Scenarios",
            ["SCN-ID", "场景描述", "速度/道路/天气等条件", "驾驶员状态", "备注"],
            c.get("operating_scenarios", []) or [],
            {"SCN-ID": "id", "场景描述": "description", "速度/道路/天气等条件": "conditions", "驾驶员状态": "driver_state", "备注": "note"},
        )
        add_table_sheet(
            "Deps",
            ["DEP-ID", "依赖项", "约束条件", "失效影响"],
            c.get("dependencies_constraints", []) or [],
            {"DEP-ID": "id", "依赖项": "dependency", "约束条件": "constraint", "失效影响": "impact_if_failed"},
        )
        add_table_sheet(
            "Concerns",
            ["CON-ID", "关注点描述", "来源", "严重性初判", "后续流转"],
            c.get("known_failures_initial_concerns", []) or [],
            {"CON-ID": "id", "关注点描述": "concern", "来源": "source", "严重性初判": "severity_preliminary", "后续流转": "downstream"},
        )
    elif meta.get("work_product_type") == "hara":
        c = source.get("content", {})
        scope = (c.get("analysis_scope", {}) or {})
        ws = wb.create_sheet(title="Scope")
        ws.append(["Field", "Value"])
        ws.append(["Related Item", scope.get("related_item", "")])
        ws.append(["Included Boundary", "; ".join(scope.get("included_boundary", []) or [])])
        ws.append(["Excluded Boundary", "; ".join(scope.get("excluded_boundary", []) or [])])

        add_table_sheet(
            "HazardEvents",
            ["HE-ID", "运行场景", "malfunction（功能异常）", "危害事件", "潜在后果"],
            c.get("hazards_and_events", []) or [],
            {
                "HE-ID": "he_id",
                "运行场景": "scenario",
                "malfunction（功能异常）": "malfunction",
                "危害事件": "hazard_event",
                "潜在后果": "consequence",
            },
        )
        add_table_sheet(
            "Risk",
            ["HE-ID", "S", "E", "C", "ASIL", "判定依据简述"],
            c.get("risk_assessment", []) or [],
            {"HE-ID": "he_id", "S": "s", "E": "e", "C": "c", "ASIL": "asil", "判定依据简述": "rationale"},
        )
        add_table_sheet(
            "SafetyGoals",
            ["SG-ID", "安全目标描述", "来源 HE-ID", "ASIL", "安全状态（Safe State）"],
            c.get("safety_goals", []) or [],
            {"SG-ID": "sg_id", "安全目标描述": "description", "来源 HE-ID": "source_he_id", "ASIL": "asil", "安全状态（Safe State）": "safe_state"},
        )
        add_table_sheet(
            "AssumptionsLimits",
            ["类型", "ID", "内容", "影响分析"],
            c.get("assumptions_limits", []) or [],
            {"类型": "type", "ID": "id", "内容": "content", "影响分析": "impact"},
        )
        add_table_sheet(
            "Downstream",
            ["SG-ID", "分配目标（FSC/FSR）", "接收方", "验证方法", "证据 ID"],
            c.get("downstream_allocation", []) or [],
            {"SG-ID": "sg_id", "分配目标（FSC/FSR）": "allocation", "接收方": "receiver", "验证方法": "verification_method", "证据 ID": "evidence_id"},
        )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    # Auto-fit columns for all sheets (except Cover), if enabled
    wb_cfg = excel_styles.get("workbook", {}) if isinstance(excel_styles, dict) else {}
    af_cfg = wb_cfg.get("column_auto_fit", {}) if isinstance(wb_cfg, dict) else {}
    for ws in wb.worksheets:
        auto_fit_sheet_columns(ws, af_cfg if isinstance(af_cfg, dict) else {})
    wb.save(out_path)


def render_docx(
    source: Dict[str, Any],
    out_path: Path,
    diagrams_by_placement: Optional[Dict[str, List[DiagramEmbed]]] = None,
) -> None:
    source = normalize_base_for_templates(source)
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception:  # pragma: no cover
        _require("python-docx")

    meta = source.get("meta", {})
    content = source.get("content", {})
    _dmap = diagrams_by_placement or {}
    _arch_dx = _dmap.get("item_architecture", [])
    _main_end_dx = _dmap.get("main_end", [])
    base_spec = load_yaml(BASE_WORD_TEMPLATE)
    word_styles = load_optional_yaml(WORD_STYLES_CONFIG)

    template_path = None
    if isinstance(word_styles, dict):
        rel = word_styles.get("docx_template_path")
        if isinstance(rel, str) and rel.strip():
            candidate = ROOT / rel
            if candidate.exists():
                template_path = candidate

    doc = Document(str(template_path)) if template_path else Document()
    title_para = doc.add_heading(meta.get("title", ""), level=0)
    if isinstance(word_styles, dict):
        title_cfg = (
            word_styles.get("styles", {}).get("title", {})
            if isinstance(word_styles.get("styles", {}), dict)
            else {}
        )
        if isinstance(title_cfg, dict):
            for run in title_para.runs:
                if "font_name" in title_cfg:
                    run.font.name = str(title_cfg["font_name"])
                if "font_size" in title_cfg:
                    run.font.size = Pt(float(title_cfg["font_size"]))
                if "bold" in title_cfg:
                    run.font.bold = bool(title_cfg["bold"])
    if include_meta_header:
        doc.add_paragraph(f"doc_id: {meta.get('doc_id','')}")
        doc.add_paragraph(f"work_product_type: {meta.get('work_product_type','')}")
        doc.add_paragraph(
            f"status: {meta.get('status','')}, version: {meta.get('version','')}, last_updated: {meta.get('last_updated','')}"
        )

    def apply_word_style_baseline() -> None:
        if not isinstance(word_styles, dict):
            return
        s = word_styles.get("styles", {}) if isinstance(word_styles.get("styles", {}), dict) else {}
        try:
            normal = doc.styles["Normal"]
            normal_cfg = s.get("normal", {}) if isinstance(s.get("normal", {}), dict) else {}
            if "font_name" in normal_cfg:
                normal.font.name = str(normal_cfg["font_name"])
            if "font_size" in normal_cfg:
                normal.font.size = Pt(float(normal_cfg["font_size"]))
            if "bold" in normal_cfg:
                normal.font.bold = bool(normal_cfg["bold"])
        except Exception:
            pass

    apply_word_style_baseline()

    def add_heading(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        if not isinstance(word_styles, dict):
            return
        s = word_styles.get("styles", {}) if isinstance(word_styles.get("styles", {}), dict) else {}
        key = "heading1" if level == 1 else ("heading2" if level == 2 else "")
        cfg = s.get(key, {}) if key and isinstance(s.get(key, {}), dict) else {}
        if not cfg:
            return
        for run in h.runs:
            if "font_name" in cfg:
                run.font.name = str(cfg["font_name"])
            if "font_size" in cfg:
                run.font.size = Pt(float(cfg["font_size"]))
            if "bold" in cfg:
                run.font.bold = bool(cfg["bold"])

    def add_bullets(items: List[Any]) -> None:
        for it in items:
            doc.add_paragraph(str(it), style="List Bullet")

    def add_table(columns: List[str], rows: List[Dict[str, Any]], mapping: Dict[str, str]) -> None:
        table = doc.add_table(rows=1, cols=len(columns))
        if isinstance(word_styles, dict):
            tbl = word_styles.get("table", {}) if isinstance(word_styles.get("table", {}), dict) else {}
            style_name = tbl.get("style_name")
            if isinstance(style_name, str) and style_name.strip():
                try:
                    table.style = style_name
                except Exception:
                    pass
        hdr = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr[i].text = col
        for r in rows:
            cells = table.add_row().cells
            for i, col in enumerate(columns):
                key = mapping[col]
                v = r.get(key, "")
                if isinstance(v, list):
                    v = ", ".join(str(x) for x in v)
                cells[i].text = str(v)

    def add_optional_docx(val: Any) -> None:
        if val is None or val == "":
            doc.add_paragraph("（待填写）")
            return
        if isinstance(val, list):
            xs = [str(x).strip() for x in val if str(x).strip()]
            if xs:
                add_bullets(xs)
            else:
                doc.add_paragraph("（待填写）")
        else:
            doc.add_paragraph(str(val))

    def add_diagram_figures(diagrams: List[DiagramEmbed]) -> None:
        if not diagrams:
            return
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            from docx.shared import Inches, Pt  # type: ignore
        except Exception:
            WD_ALIGN_PARAGRAPH = None  # type: ignore[misc, assignment]
            Inches = None  # type: ignore[misc, assignment]
            Pt = None  # type: ignore[misc, assignment]

        # A4 正文区约 6.0–6.2"；限制最大宽/高，避免单图撑满页或超出可视区域
        max_w_in = 5.5
        max_h_in = 4.25
        assume_dpi = 96.0

        for d in diagrams:
            cap = doc.add_paragraph(str(d.title))
            if WD_ALIGN_PARAGRAPH is not None:
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if Pt is not None:
                cap.paragraph_format.space_after = Pt(6)

            if Inches is None or not d.png_path.is_file():
                doc.add_paragraph(f"（图未渲染或缺少 PNG；Mermaid 源：{d.rel_mmd}）")
                continue

            px = _png_ihdr_pixel_size(d.png_path)
            if px:
                w_px, h_px = px
                w_in = w_px / assume_dpi
                h_in = h_px / assume_dpi
                scale = min(max_w_in / max(w_in, 1e-6), max_h_in / max(h_in, 1e-6), 1.0)
                out_w_in = w_in * scale
                width = Inches(out_w_in)
            else:
                width = Inches(max_w_in)

            pic_p = doc.add_paragraph()
            if WD_ALIGN_PARAGRAPH is not None:
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_p.add_run()
            run.add_picture(str(d.png_path), width=width)
            if Pt is not None:
                pic_p.paragraph_format.space_after = Pt(12)

    def append_item_definition_docx_main() -> None:
        c = content if isinstance(content, dict) else {}
        ov = c.get("item_overview", {}) or {}
        bi = c.get("boundary_interfaces", {}) or {}
        da = c.get("driving_automation", {}) or {}
        odd_raw = c.get("odd")
        vss = c.get("vehicle_safety_strategy", {}) or {}
        ad = c.get("alerts_degradation", {}) or {}
        fp = c.get("function_partition")
        frs = c.get("functional_runtime_state")
        fm = c.get("foreseeable_misuse")
        nf = c.get("non_functional", {}) or {}
        dep_o2t = c.get("dependency_other_to_this")
        dep_t2o = c.get("dependency_this_to_other")
        dep_env = c.get("dependency_environment")
        kfm = c.get("known_failure_modes")
        ktc = c.get("known_trigger_conditions")
        osa = c.get("operational_safety_assurance", {}) or {}
        impl = c.get("implementation", {}) or {}

        add_heading("相关项定义", level=1)
        add_heading("功能定义", level=2)
        add_heading("功能简介", level=3)
        doc.add_paragraph(f"Item 名称：{ov.get('item_name', '')}")
        doc.add_paragraph(f"功能目标：{ov.get('functional_objective', '')}")
        doc.add_paragraph("非安全目标（如性能/舒适性）：")
        nso = ov.get("non_safety_objectives", []) or []
        if nso:
            add_bullets(nso)
        else:
            doc.add_paragraph("（待填写）")

        add_heading("驾驶自动化等级", level=3)
        if isinstance(da, dict) and (da.get("level") or da.get("standard") or da.get("role_phases")):
            if da.get("level"):
                doc.add_paragraph(f"驾驶自动化等级：{da.get('level', '')}")
            if da.get("standard"):
                doc.add_paragraph(f"参考标准：{da.get('standard', '')}")
            rp = da.get("role_phases") or []
            if rp:
                add_table(
                    ["阶段", "用户的角色", "系统的角色"],
                    rp,
                    {"阶段": "phase", "用户的角色": "user_role", "系统的角色": "system_role"},
                )
        else:
            doc.add_paragraph("（待填写；非驾驶自动化相关项无需填写本节。）")

        add_heading("ODD（设计运行范围）", level=3)
        if isinstance(odd_raw, dict):
            add_optional_docx(odd_raw.get("reference") or odd_raw.get("description"))
        else:
            add_optional_docx(odd_raw)

        add_heading("整车层面安全策略", level=3)
        add_heading("整车安全策略", level=4)
        add_heading("法律法规", level=5)
        add_optional_docx(vss.get("laws"))
        add_heading("驾驶行为", level=5)
        add_optional_docx(vss.get("driving_behavior"))

        add_heading("整车层面 SOTIF 残余风险接受准则", level=4)
        add_optional_docx(vss.get("sotif_residual"))

        add_heading("提示报警与降级", level=4)
        add_heading("提示报警", level=5)
        add_optional_docx(ad.get("alerts") if ad else vss.get("alerts"))
        add_heading("降级思路", level=5)
        add_optional_docx(ad.get("rationale") if ad else vss.get("degradation_rationale"))
        add_heading("降级策略", level=5)
        rm_states = (ad.get("rm_states") if ad else None) or vss.get("rm_states")
        triggers = (ad.get("triggers") if ad else None) or vss.get("triggers")
        if rm_states:
            doc.add_paragraph("状态说明")
            add_table(
                ["状态", "说明", "备注"],
                rm_states,
                {"状态": "state", "说明": "description", "备注": "note"},
            )
        if triggers:
            doc.add_paragraph("跳转条件说明")
            add_table(
                ["跳转条件", "说明", "备注"],
                triggers,
                {"跳转条件": "trigger", "说明": "description", "备注": "note"},
            )
        if not rm_states and not triggers:
            doc.add_paragraph("（待填写）")

        add_heading("功能行为与用例", level=2)
        add_heading("功能划分与简介", level=3)
        if isinstance(fp, list) and fp and isinstance(fp[0], dict):
            add_table(
                ["子功能", "说明", "备注"],
                fp,
                {"子功能": "subfunction", "说明": "description", "备注": "note"},
            )
        elif fp:
            add_bullets(ensure_list(fp))
        else:
            doc.add_paragraph("（可与「功能简介」「系统边界」一致；以下为当前边界中的在范围要点）")
            ins = bi.get("in_scope", []) or []
            if ins:
                add_bullets(ins[:12])
            else:
                doc.add_paragraph("（待填写）")

        add_heading("用例说明", level=3)
        ucs = c.get("use_cases")
        if isinstance(ucs, list) and ucs:
            for i, u in enumerate(ucs):
                if not isinstance(u, dict):
                    continue
                uid = str(u.get("id", f"UC-{i + 1}"))
                name = str(u.get("name", "")).strip()
                add_heading(f"用例 {uid}" + (f"：{name}" if name else ""), level=4)
                rows = [
                    {"k": "用例名称", "v": name},
                    {"k": "主要参与者", "v": u.get("primary_actors", "")},
                    {"k": "次要参与者", "v": u.get("secondary_actors", "")},
                    {"k": "前置条件", "v": u.get("preconditions", "")},
                    {"k": "主成功场景", "v": u.get("main_success", "")},
                    {"k": "拓展场景", "v": u.get("extensions", "")},
                ]
                add_table(
                    ["字段", "内容"],
                    rows,
                    {"字段": "k", "内容": "v"},
                )
        else:
            add_table(
                ["SCN-ID", "场景描述", "速度/道路/天气等条件", "驾驶员状态", "备注"],
                c.get("operating_scenarios", []) or [],
                {
                    "SCN-ID": "id",
                    "场景描述": "description",
                    "速度/道路/天气等条件": "conditions",
                    "驾驶员状态": "driver_state",
                    "备注": "note",
                },
            )

        add_heading("功能运行状态", level=2)
        if isinstance(frs, dict):
            st = frs.get("states")
            tr = frs.get("transitions")
            if isinstance(st, list) and st:
                add_table(
                    ["状态", "说明", "备注"],
                    st,
                    {"状态": "state", "说明": "description", "备注": "note"},
                )
            if isinstance(tr, list) and tr:
                add_table(
                    ["跳转条件", "说明", "备注"],
                    tr,
                    {"跳转条件": "trigger", "说明": "description", "备注": "note"},
                )
            extra = frs.get("notes") or frs.get("narrative")
            if extra:
                add_optional_docx(extra)
            if not st and not tr and not extra:
                doc.add_paragraph("（待填写）")
        else:
            add_optional_docx(frs)

        add_heading("合理可预见的误用", level=2)
        if isinstance(fm, list) and fm and isinstance(fm[0], dict):
            add_table(
                ["编号", "误用说明", "备注"],
                fm,
                {"编号": "id", "误用说明": "description", "备注": "note"},
            )
        else:
            add_optional_docx(fm)

        add_heading("非功能性要求", level=2)
        add_heading("质量要求", level=3)
        add_optional_docx(nf.get("quality"))
        add_heading("性能与可用性要求", level=3)
        add_optional_docx(nf.get("performance"))

        add_heading("功能依赖 / 假设", level=2)
        add_heading("其他相关项对本相关项的依赖 / 假设", level=3)
        add_optional_docx(dep_o2t)
        add_heading("本相关项对其他相关项的依赖 / 假设", level=3)
        add_optional_docx(dep_t2o)
        add_heading("环境依赖 / 假设", level=3)
        add_optional_docx(dep_env)
        doc.add_paragraph("依赖与约束（结构化表）")
        add_table(
            ["DEP-ID", "依赖项（系统/传感器/电源/网络）", "约束条件", "失效影响"],
            c.get("dependencies_constraints", []) or [],
            {
                "DEP-ID": "id",
                "依赖项（系统/传感器/电源/网络）": "dependency",
                "约束条件": "constraint",
                "失效影响": "impact_if_failed",
            },
        )

        add_heading("功能层面已知功能局限", level=2)
        add_heading("已知失效模式", level=3)
        add_optional_docx(kfm)
        add_heading("已知功能触发条件", level=3)
        add_optional_docx(ktc)
        doc.add_paragraph("初始安全关注点（供 HARA 输入）")
        add_table(
            ["CON-ID", "关注点描述", "来源", "严重性初判", "后续流转"],
            c.get("known_failures_initial_concerns", []) or [],
            {
                "CON-ID": "id",
                "关注点描述": "concern",
                "来源": "source",
                "严重性初判": "severity_preliminary",
                "后续流转": "downstream",
            },
        )

        add_heading("运行安全保障", level=2)
        add_heading("运行阶段安全设计", level=3)
        if isinstance(osa, dict):
            add_optional_docx(osa.get("operational_design"))
        else:
            add_optional_docx(osa)

        add_heading("功能实现", level=2)
        add_heading("功能架构", level=3)
        add_heading("架构示意图", level=4)
        add_optional_docx(impl.get("architecture"))
        add_diagram_figures(_arch_dx)
        add_heading("要素说明", level=4)
        arch_el = impl.get("architecture_elements")
        if isinstance(arch_el, list) and arch_el:
            add_table(
                ["要素", "说明", "备注"],
                arch_el,
                {"要素": "element", "说明": "description", "备注": "note"},
            )
        else:
            doc.add_paragraph("在范围（In Scope）：")
            add_bullets(bi.get("in_scope", []) or ["（待填写）"])
            doc.add_paragraph("不在范围（Out of Scope）：")
            add_bullets(bi.get("out_of_scope", []) or ["（待填写）"])
        add_heading("接口说明", level=4)
        ifb = impl.get("interfaces_brief")
        if isinstance(ifb, list) and ifb:
            add_table(
                ["接口", "说明", "备注"],
                ifb,
                {"接口": "interface", "说明": "description", "备注": "note"},
            )
        else:
            add_table(
                ["IF-ID", "接口对象", "输入/输出", "信号/消息", "正常范围", "异常处理"],
                bi.get("interfaces", []) or [],
                {
                    "IF-ID": "id",
                    "接口对象": "object",
                    "输入/输出": "direction",
                    "信号/消息": "signal",
                    "正常范围": "normal_range",
                    "异常处理": "abnormal_handling",
                },
            )

        add_heading("非功能性要求（实现层）", level=3)
        doc.add_paragraph("需要从概念阶段的非功能性要求进行拆解得出。")
        if impl.get("nf_quality"):
            add_heading("质量要求", level=4)
            add_optional_docx(impl.get("nf_quality"))
        add_heading("性能与可用性要求", level=4)
        add_optional_docx(impl.get("nf_performance"))

        add_heading("要素层面已知功能局限", level=3)
        add_heading("已知失效模式", level=4)
        add_optional_docx(impl.get("element_failure_modes"))
        add_heading("已知功能触发条件", level=4)
        add_optional_docx(impl.get("element_trigger_conditions"))

    # Render base sections driven by template spec（占位节不写标题，避免与正文重复出现「正文内容」）
    for sec in (base_spec or {}).get("sections", []) or []:
        if sec.get("content_placeholder"):
            break
        heading = sec.get("heading")
        if heading:
            hl = int(sec.get("heading_level", 1))
            add_heading(str(heading), level=hl)
        if "source" in sec:
            v = get_path(source, str(sec.get("source")))
            if isinstance(v, list):
                add_bullets(v)
            elif v is not None:
                doc.add_paragraph(str(v))
        if "table" in sec:
            t = sec.get("table", {}) or {}
            cols = t.get("columns", []) or []
            rows = get_path(source, str(t.get("source"))) if t.get("source") else []
            if rows is None:
                rows = []
            add_table(cols, rows or [], t.get("mapping", {}) or {})
        if "blocks" in sec:
            for blk in sec.get("blocks", []) or []:
                sub = blk.get("subheading")
                if sub:
                    add_heading(str(sub), level=2)
                if "source" in blk:
                    v = get_path(source, str(blk.get("source")))
                    if blk.get("list_style") == "bullet":
                        add_bullets(v if isinstance(v, list) else ensure_list(v))
                    else:
                        if v is not None:
                            doc.add_paragraph(str(v))
                if "table" in blk:
                    t = blk.get("table", {}) or {}
                    cols = t.get("columns", []) or []
                    rows = get_path(source, str(t.get("source"))) if t.get("source") else []
                    if rows is None:
                        rows = []
                    add_table(cols, rows or [], t.get("mapping", {}) or {})

    if meta.get("work_product_type") == "item-definition":
        append_item_definition_docx_main()
    elif meta.get("work_product_type") == "hara":
        c = source.get("content", {})
        scope = (c.get("analysis_scope", {}) or {})
        doc.add_heading("危害分析与风险评估（HARA）", level=1)
        doc.add_heading("分析范围", level=2)
        doc.add_paragraph(f"关联 Item：{scope.get('related_item','')}")
        doc.add_paragraph("分析边界：")
        add_bullets(scope.get("included_boundary", []) or [])
        doc.add_paragraph("不覆盖范围：")
        add_bullets(scope.get("excluded_boundary", []) or [])

        doc.add_heading("场景与危害事件识别", level=2)
        add_table(
            ["HE-ID", "运行场景", "malfunction（功能异常）", "危害事件", "潜在后果"],
            c.get("hazards_and_events", []) or [],
            {"HE-ID": "he_id", "运行场景": "scenario", "malfunction（功能异常）": "malfunction", "危害事件": "hazard_event", "潜在后果": "consequence"},
        )
        doc.add_heading("风险评估与 ASIL 判定", level=2)
        add_table(
            ["HE-ID", "S", "E", "C", "ASIL", "判定依据简述"],
            c.get("risk_assessment", []) or [],
            {"HE-ID": "he_id", "S": "s", "E": "e", "C": "c", "ASIL": "asil", "判定依据简述": "rationale"},
        )
        doc.add_heading("安全目标（Safety Goals）", level=2)
        add_table(
            ["SG-ID", "安全目标描述", "来源 HE-ID", "ASIL", "安全状态（Safe State）"],
            c.get("safety_goals", []) or [],
            {"SG-ID": "sg_id", "安全目标描述": "description", "来源 HE-ID": "source_he_id", "ASIL": "asil", "安全状态（Safe State）": "safe_state"},
        )
        doc.add_heading("假设、限制与边界条件", level=2)
        add_table(
            ["类型", "ID", "内容", "影响分析"],
            c.get("assumptions_limits", []) or [],
            {"类型": "type", "ID": "id", "内容": "content", "影响分析": "impact"},
        )
        doc.add_heading("向下游分配与追溯", level=2)
        add_table(
            ["SG-ID", "分配目标（FSC/FSR）", "接收方", "验证方法", "证据 ID"],
            c.get("downstream_allocation", []) or [],
            {"SG-ID": "sg_id", "分配目标（FSC/FSR）": "allocation", "接收方": "receiver", "验证方法": "verification_method", "证据 ID": "evidence_id"},
        )
        if _main_end_dx:
            doc.add_heading("UML 附图", level=2)
            add_diagram_figures(_main_end_dx)
    else:
        doc.add_paragraph("当前仅实现 item-definition 与 hara 的正文渲染。")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    p = argparse.ArgumentParser(description="Render MD/DOCX/XLSX from YAML single source of truth.")
    p.add_argument("--input", required=True, help="YAML input path")
    p.add_argument("--out-dir", default=str(ROOT / "out"), help="Output directory")
    p.add_argument(
        "--formats",
        default=None,
        help="Comma-separated output formats. If omitted, defaults are selected by work_product_type.",
    )
    p.add_argument(
        "--include-meta",
        action="store_true",
        help="Include meta header blocks in generated documents (MD front matter, DOCX meta paragraphs, XLSX cover sheet).",
    )
    p.add_argument(
        "--flat-output",
        action="store_true",
        help="将 MD/DOCX/XLSX 直接写入 --out-dir 根目录（旧行为）；默认写入 --out-dir/<doc_id>/ 并附带 uml/ 子目录。",
    )
    args = p.parse_args()

    src_path = Path(args.input).resolve()
    data = load_yaml(src_path)
    meta = data.get("meta", {})
    wp_type = str(meta.get("work_product_type", "")).strip()
    if not wp_type:
        raise RuntimeError("YAML 缺少 meta.work_product_type，无法选择默认输出格式")
    doc_id = meta.get("doc_id", src_path.stem)
    out_dir = Path(args.out_dir).resolve()
    formats = set(resolve_formats(wp_type, args.formats))

    if args.flat_output:
        md_parent = out_dir
        uml_dir = out_dir / "uml" / str(doc_id)
    else:
        md_parent = out_dir / str(doc_id)
        uml_dir = md_parent / "uml"
    md_parent.mkdir(parents=True, exist_ok=True)

    diagram_map = prepare_diagram_embeds(data, uml_dir, md_parent)

    if "md" in formats:
        md = render_md(data, diagram_map)
        if args.include_meta:
            fm_lines = ["---"]
            for k in [
                "doc_id",
                "title",
                "work_product_type",
                "iso_ref",
                "asil",
                "status",
                "owner",
                "reviewers",
                "version",
                "baseline",
                "last_updated",
            ]:
                v = meta.get(k)
                if isinstance(v, list):
                    fm_lines.append(f"{k}:")
                    for item in v:
                        fm_lines.append(f"  - {item}")
                else:
                    fm_lines.append(f"{k}: {v}")
            fm_lines.append("---\n")
            md = "\n".join(fm_lines) + md
        md_path = md_parent / f"{doc_id}.md"
        md_path.write_text(md, encoding="utf-8")
        print(f"[OK] MD -> {md_path}")
    if "xlsx" in formats:
        xlsx_path = md_parent / f"{doc_id}.xlsx"
        render_xlsx(data, xlsx_path)
        print(f"[OK] XLSX -> {xlsx_path}")
    if "docx" in formats:
        docx_path = md_parent / f"{doc_id}.docx"
        global include_meta_header
        include_meta_header = bool(args.include_meta)
        render_docx(data, docx_path, diagram_map)
        print(f"[OK] DOCX -> {docx_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

